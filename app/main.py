from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pathlib import Path
import shutil, uuid, json
from typing import List

from app.schemas.models import OneLineRequest, PlanRequest
from app.cad.one_line import generate_one_line_dxf
from app.cad.power_plan import generate_power_plan_dxf
from app.cad.lighting_plan import generate_lighting_plan_dxf
from app.ai.llm import plan_from_prompt, summarize_intent, extract_panel_specs_from_text
from app.db import init_db, get_active_task, save_task_state, update_task_parameters, clear_task_state
from app.utils.excel_template import find_template, extract_template_parameters
from app.routers import panel as panel_router
from app.routers import preflight


ROOT = Path(__file__).resolve().parent.parent
BUCKET = ROOT / "bucket"
OUT = ROOT / "out"
STATIC = ROOT / "static"

app = FastAPI(title="AI Design Engineer Voice & Text Assistant")

# Register panel preflight checks
app.include_router(preflight.router)
# Register panel schedule routes (XLSX/PDF export)
app.include_router(panel_router.router)

 
# CORS (single block)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
def startup_event():
    init_db()

# Static frontend
app.mount("/static", StaticFiles(directory=str(STATIC), html=True), name="static")

@app.get("/", response_class=HTMLResponse)
def home():
    index = (STATIC / "index.html").read_text(encoding="utf-8")
    return HTMLResponse(index)

@app.get("/health")
def health():
    return {"status": "ok"}

# ---- Bucket (drag & drop) ----
@app.post("/bucket/upload")
async def upload(files: List[UploadFile] = File(...), session: str | None = None):
    BUCKET.mkdir(parents=True, exist_ok=True)
    saved = []
    template_detected = False
    pref = _session_prefix(session)
    for f in files:
        filename = pref + f.filename
        dest = BUCKET / filename
        with dest.open("wb") as out:
            shutil.copyfileobj(f.file, out)
        saved.append(filename)
        
        # Detect if this is a template file
        if 'template' in f.filename.lower() and f.filename.lower().endswith(('.xlsx', '.xlsm')):
            template_detected = True
    
    response = {"saved": saved}
    if template_detected:
        response["template_detected"] = True
        response["message"] = "A template file has been uploaded. This template will be used for panel schedules in this session. The default template remains unchanged for future sessions."
    
    return response

@app.get("/bucket/list")
def bucket_list(session: str | None = None):
    BUCKET.mkdir(exist_ok=True)
    files = sorted([p.name for p in BUCKET.iterdir() if p.is_file()])
    return {"files": _filter_session(files, session)}

@app.get("/bucket/file/{name}")
def bucket_file(name: str):
    path = BUCKET / name
    if not path.exists():
        raise HTTPException(404, "Not found")
    return FileResponse(str(path))

@app.post("/bucket/clear")
def bucket_clear(session: str | None = None):
    if BUCKET.exists():
        pref = _session_prefix(session)
        for p in list(BUCKET.iterdir()):
            if p.is_file() and (not pref or p.name.startswith(pref)):
                p.unlink()
    return {"status": "cleared"}

# ---- Outputs list (for UI) ----
@app.get("/outputs/list")
def outputs_list(session: str | None = None):
    OUT.mkdir(exist_ok=True)
    files = sorted([p.name for p in OUT.iterdir() if p.is_file()])
    return {"files": _filter_session(files, session)}

@app.get("/out/{name}")
def out_file(name: str):
    path = OUT / name
    if not path.exists():
        raise HTTPException(404, "Not found")
    return FileResponse(str(path))

# ---- CAD endpoints (unchanged programmatic access) ----
@app.post("/cad/one_line")
def cad_one_line(req: OneLineRequest):
    OUT.mkdir(parents=True, exist_ok=True)
    out_path = OUT / _short_filename('one_line', 'dxf')
    generate_one_line_dxf(req, out_path)
    return {"file": out_path.name}

@app.post("/cad/power_plan")
def cad_power_plan(req: PlanRequest):
    OUT.mkdir(parents=True, exist_ok=True)
    out_path = OUT / _short_filename('power_plan', 'dxf')
    generate_power_plan_dxf(req, out_path)
    return {"file": out_path.name}

@app.post("/cad/lighting_plan")
def cad_lighting_plan(req: PlanRequest):
    OUT.mkdir(parents=True, exist_ok=True)
    out_path = OUT / _short_filename('lighting_plan', 'dxf')
    generate_lighting_plan_dxf(req, out_path)
    return {"file": out_path.name}

# ---- Voice/typed command dispatcher ----

def _filter_plan_params(params: dict) -> dict:
    """Remove internal state flags from parameters before including in plan response."""
    return {k: v for k, v in params.items() if k not in ["pending_confirmation", "pending_finish", "prompt_count"]}

@app.post("/commands/run")
def run_command(payload: dict):
    text = (payload.get("text") or "").strip()
    if not text:
        raise HTTPException(400, "No command text provided.")

    session = payload.get("session")
    if not session:
        raise HTTPException(400, "Session ID required.")
    
    pref = _session_prefix(session)
    
    # Check for reference files in bucket
    BUCKET.mkdir(exist_ok=True)
    bucket_files = [p.name for p in BUCKET.iterdir() if p.is_file()]
    session_files = [f for f in bucket_files if not pref or f.startswith(pref)]
    
    # Check for Yes/No responses
    text_lower = text.lower().strip()
    if text_lower in ["yes", "y", "yeah", "yep", "sure", "ok", "okay"]:
        # User confirmed something - check what we're confirming
        active_task = get_active_task(session)
        
        # IMPORTANT: Check pending_finish FIRST (takes priority over pending_confirmation)
        if active_task and active_task.get("parameters", {}).get("pending_finish"):
            # User confirmed finishing the task
            task_name = active_task["task_type"].replace("_", " ")
            clear_task_state(session)
            return {
                "summary": "Got it.",
                "message": f"{task_name.title()} finished.",
                "plan": {"task": "none", "project": "Ready"}
            }
        elif active_task and active_task.get("parameters", {}).get("pending_confirmation"):
            # User is confirming to start the task
            params = active_task["parameters"]
            params.pop("pending_confirmation", None)
            update_task_parameters(session, params)
            
            task_type = active_task["task_type"]
            task_name = task_type.replace("_", " ")
            
            # For panel_schedule, extract template parameters and check requirements
            if task_type == "panel_schedule":
                # Extract template parameters if not already done
                if "template_parameters" not in params:
                    template_path = find_template(BUCKET, pref)
                    if template_path:
                        template_params = extract_template_parameters(template_path)
                        params["template_parameters"] = template_params
                        update_task_parameters(session, params)
                
                number_of_ckts = params.get("number_of_ckts")
                if not number_of_ckts:
                    return {
                        "summary": "Got it.",
                        "message": "How many circuits? (Please provide an even number between 18-84)",
                        "plan": {"task": task_type, "project": params.get("project", "Project"), **_filter_plan_params(params)},
                        "needs_input": "number_of_ckts"
                    }
                else:
                    file_info = ""
                    if session_files:
                        file_info = f" I see {len(session_files)} reference file(s)."
                    else:
                        file_info = " Upload panel photos for better results."
                    
                    return {
                        "summary": "Got it.",
                        "message": f"Ready to build {number_of_ckts}-circuit panel schedule.{file_info} Press Build when ready. (Say 'finished' when done)",
                        "plan": {"task": task_type, "project": params.get("project", "Project"), **_filter_plan_params(params)}
                    }
            
            # For other tasks, ready to go
            return {
                "summary": "Got it.",
                "message": f"Starting work on {task_name}. Press Build when ready. (Say 'finished' when done)",
                "plan": {"task": task_type, "project": params.get("project", "Project"), **_filter_plan_params(params)}
            }
        else:
            # Generic yes without context - treat as continue
            pass
    
    if text_lower in ["no", "n", "nope", "nah", "cancel"]:
        # User declined something - check what we're declining
        active_task = get_active_task(session)
        if active_task and active_task.get("parameters", {}).get("pending_confirmation"):
            # User declined to start the task
            clear_task_state(session)
            return {
                "summary": "Got it.",
                "message": "No problem. What would you like to do instead?",
                "plan": {"task": "none", "project": "Ready"}
            }
        elif active_task and active_task.get("parameters", {}).get("pending_finish"):
            # User declined to finish the task - stay on it
            params = active_task["parameters"]
            params.pop("pending_finish", None)
            update_task_parameters(session, params)
            
            task_type = active_task["task_type"]
            task_name = task_type.replace("_", " ")
            
            return {
                "summary": "Got it.",
                "message": f"Continuing work on {task_name}. What else do you need?",
                "plan": {"task": task_type, "project": params.get("project", "Project"), **_filter_plan_params(params)}
            }
    
    # Check if user wants to finish the current task
    finish_keywords = ["finished", "done", "complete", "stop"]
    if any(kw in text_lower for kw in finish_keywords):
        active_task = get_active_task(session)
        if active_task:
            task_type = active_task["task_type"]
            task_name = task_type.replace("_", " ")
            
            # Ask for confirmation
            params = active_task["parameters"]
            params["pending_finish"] = True
            update_task_parameters(session, params)
            
            return {
                "summary": "Got it.",
                "message": f"Do you want to finish the {task_name}?",
                "plan": {"task": task_type, "project": params.get("project", "Project"), **_filter_plan_params(params)},
                "needs_finish_confirmation": True,
                "task_name": task_name
            }
        else:
            return {
                "summary": "Got it.",
                "message": "No active task to finish. What would you like to do?",
                "plan": {"task": "none", "project": "Ready"}
            }
    
    # Check for active task
    active_task = get_active_task(session)
    
    if active_task:
        # User is continuing an active task, extract parameters from their response
        task_type = active_task["task_type"]
        params = active_task["parameters"]
        
        # Parse the user's response to extract parameters
        new_plan = plan_from_prompt(text, str(BUCKET))
        
        # Track which parameters were newly extracted or updated
        extracted_params = []
        
        # Extract panel specs from text (voltage, phase, etc.)
        panel_specs_from_text = extract_panel_specs_from_text(text)
        if panel_specs_from_text:
            if "panel_specs" not in params:
                params["panel_specs"] = {}
            params["panel_specs"].update(panel_specs_from_text)
            for key, value in panel_specs_from_text.items():
                extracted_params.append(f"{key} is {value}")
        
        # Update parameters from user response (allow overrides/corrections)
        if new_plan.get("number_of_ckts"):
            old_value = params.get("number_of_ckts")
            new_value = new_plan["number_of_ckts"]
            if old_value != new_value:
                params["number_of_ckts"] = new_value
                extracted_params.append(f"number of circuits is {new_value}")
        if new_plan.get("panel_name"):
            old_value = params.get("panel_name")
            new_value = new_plan["panel_name"]
            if old_value != new_value:
                params["panel_name"] = new_value
                extracted_params.append(f"panel name is {new_value}")
        
        # Save updated parameters
        update_task_parameters(session, params)
        
        # Build the plan with updated parameters (exclude internal state flags)
        plan = {
            "task": task_type,
            "project": params.get("project", "Project"),
            **_filter_plan_params(params)
        }
        
        # Build confirmation message
        confirmation = "Got it."
        if extracted_params:
            confirmation = f"Got it, {', '.join(extracted_params)}."
        
        # Check if we have all required parameters for panel_schedule
        if task_type == "panel_schedule":
            # Ensure template parameters are extracted
            if "template_parameters" not in params:
                template_path = find_template(BUCKET, pref)
                if template_path:
                    template_params = extract_template_parameters(template_path)
                    params["template_parameters"] = template_params
                    update_task_parameters(session, params)
            
            number_of_ckts = params.get("number_of_ckts")
            
            if not number_of_ckts:
                return {
                    "summary": confirmation,
                    "message": "How many circuits? (Please provide an even number between 18-84)",
                    "plan": plan,
                    "needs_input": "number_of_ckts"
                }
            
            # All parameters collected!
            # Track if this is the first time we have circuit count
            first_circuit_count = "number_of_ckts" not in params or params.get("number_of_ckts") != number_of_ckts
            
            # Track prompt count to avoid repetitive "Ready to build..." message
            if first_circuit_count:
                params["prompt_count"] = 1  # Reset counter when circuit count is first set/changed
            else:
                params["prompt_count"] = params.get("prompt_count", 0) + 1
            
            prompt_count = params["prompt_count"]
            update_task_parameters(session, params)
            
            # Show "Ready to build..." message when circuit count is first set, or every 10 prompts after
            if first_circuit_count or prompt_count % 10 == 0:
                file_info = ""
                if session_files:
                    file_info = f" I see {len(session_files)} reference file(s)."
                else:
                    file_info = " Upload panel photos for better results."
                
                message = f"Ready to build {number_of_ckts}-circuit panel schedule.{file_info} Press Build when ready. (Say 'finished' to end your task)"
            else:
                # Just acknowledge without repeating the ready message
                message = "Got it." if not confirmation or confirmation == "Got it." else confirmation
            
            return {
                "summary": confirmation,
                "message": message,
                "plan": plan
            }
        
        # For other tasks, return ready message
        return {
            "summary": "Got it.",
            "message": f"Ready for {task_type.replace('_', ' ')}. Press Build when ready. (Say 'finished' to end your task)",
            "plan": plan
        }
    
    # No active task, parse as new command
    summary = summarize_intent(text)
    plan = plan_from_prompt(text, str(BUCKET))
    task = (plan.get("task") or "").lower()
    
    # Check if this is a recognized task that needs confirmation
    if task in ["panel_schedule", "one_line", "power_plan", "lighting_plan", "revit_package"]:
        task_name = task.replace("_", " ")
        
        # Initialize task state with pending_confirmation flag
        params = {
            "project": plan.get("project", "Project"),
            "pending_confirmation": True
        }
        
        # Save any task-specific parameters from the initial command
        if task == "panel_schedule":
            if plan.get("number_of_ckts"):
                params["number_of_ckts"] = plan.get("number_of_ckts")
            if plan.get("panel_name"):
                params["panel_name"] = plan.get("panel_name")
        
        save_task_state(session, task, params)
        
        return {
            "summary": summary,
            "message": f"Should I work on a {task_name}?",
            "plan": plan,
            "needs_confirmation": True,
            "task_name": task_name
        }

    # Unrecognized task or general query
    return {
        "summary": summary, 
        "message": "I'm listening. What would you like to do?",
        "plan": plan
    }


from app.schemas.standards import StandardsConfig

STANDARDS_DIR = ROOT / "standards"
ACTIVE_STANDARDS = STANDARDS_DIR / "active.json"

def load_standards() -> StandardsConfig:
    STANDARDS_DIR.mkdir(exist_ok=True)
    if ACTIVE_STANDARDS.exists():
        try:
            return StandardsConfig(**json.loads(ACTIVE_STANDARDS.read_text()))
        except Exception:
            pass
    cfg = StandardsConfig()
    ACTIVE_STANDARDS.write_text(json.dumps(cfg.model_dump(), indent=2))
    return cfg

@app.post("/standards/upload")
async def standards_upload(config: UploadFile = File(None), titleblock: UploadFile = File(None)):
    STANDARDS_DIR.mkdir(exist_ok=True)
    result = {}
    if config is not None:
        text = config.file.read().decode("utf-8", errors="ignore")
        try:
            cfg_json = json.loads(text)
        except Exception:
            # fallback: simple key:value parser
            cfg_json = {}
            for line in text.splitlines():
                if ":" in line and not line.strip().startswith("#"):
                    k,v = line.split(":",1)
                    cfg_json[k.strip()] = v.strip()
            if "layers" not in cfg_json:
                cfg_json = {"layers": {"annotations": "E-ANNO-TEXT"}}
        ACTIVE_STANDARDS.write_text(json.dumps(cfg_json, indent=2))
        result["config"] = "saved"
    if titleblock is not None:
        tb_name = titleblock.filename
        tb_path = STANDARDS_DIR / tb_name
        with tb_path.open("wb") as f:
            f.write(titleblock.file.read())
        cfg = load_standards()
        cfg.titleblock = tb_name
        ACTIVE_STANDARDS.write_text(json.dumps(cfg.model_dump(), indent=2))
        result["titleblock"] = tb_name
    if not result:
        raise HTTPException(400, "No files uploaded.")
    return result

@app.get("/standards/get")
def standards_get():
    cfg = load_standards()
    return cfg.model_dump()


from fastapi import Body
import csv, io

@app.post("/cad/panel_schedule_csv")
def cad_panel_schedule_csv(req: OneLineRequest = Body(...)):
    """Export a very simple panel schedule CSV from the OneLineRequest."""
    # Aggregate loads by panel
    by_panel = {}
    for ld in req.loads:
        by_panel.setdefault(ld.panel, []).append(ld)
    csv_buf = io.StringIO()
    writer = csv.writer(csv_buf)
    writer.writerow(["Panel", "Load Name", "kVA"])
    for pnl, loads in by_panel.items():
        for ld in loads:
            writer.writerow([pnl, ld.name, ld.kva])
    content = csv_buf.getvalue().encode("utf-8")
    out_path = OUT / _short_filename('panel_schedule', 'csv')
    with open(out_path, "wb") as f:
        f.write(content)
    return {"file": out_path.name}



# ---- DXF → PDF export ----
from app.export.pdf_from_dxf import dxf_to_pdf

@app.get("/export/pdf")
def export_pdf(file: str, session: str | None = None):
    # convert an existing DXF in OUT directory to a PDF with same stem
    dxf_path = OUT / file
    if not dxf_path.exists() or not dxf_path.suffix.lower() in [".dxf"]:
        raise HTTPException(400, "Provide a DXF file name that exists in /outputs/list.")
    pdf_name = Path(file).with_suffix(".pdf").name
    pdf_path = OUT / pdf_name
    dxf_to_pdf(dxf_path, pdf_path)
    return {"message": "PDF generated.", "file": pdf_name}


# ---- Build ZIP with summary DOCX ----
from docx import Document
from datetime import datetime
import zipfile as _zipfile

def _write_summary_docx(plan: dict, out_dir: Path, session: str | None = None) -> Path:
    doc = Document()
    doc.add_heading('AI Design Engineer - Summary', level=1)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    doc.add_paragraph(f"Timestamp: {now}")
    doc.add_paragraph(" ")
    # Plan echo
    doc.add_heading('What was requested (my understanding)', level=2)
    doc.add_paragraph(json.dumps(plan, indent=2))

    # What was done
    doc.add_heading('What was done & why', level=2)
    doc.add_paragraph("Generated draft design artifacts based on your instructions and project standards. "
                      "Outputs are provided in the ZIP. Rationale: match scope to request, adhere to configured layers/symbols, "
                      "and provide deterministic CAD structure for downstream PE review.")

    # Ethics
    doc.add_heading('Ethical considerations', level=2)
    doc.add_paragraph("This draft was produced by an AI-assisted workflow. It should be reviewed by a licensed Professional Engineer "
                      "before use in construction. The intent is to improve efficiency while maintaining public safety and code compliance. "
                      "All generative steps are logged via inputs/plan to preserve auditability.")

    # PE review
    doc.add_heading('PE review checklist', level=2)
    doc.add_paragraph("- Confirm service size, feeder/breaker coordination, and fault current assumptions.\n"
                      "- Verify equipment ratings and selective coordination where required.\n"
                      "- Validate conductor sizes, voltage drop and grounding/bonding details.\n"
                      "- Confirm device locations against architectural constraints and egress.\n"
                      "- Stamp only after verifying that calculations, notes, and details meet local/state amendments.")

    # NEC references (generic, non-exhaustive)
    doc.add_heading('Relevant NEC sections (non-exhaustive)', level=2)
    doc.add_paragraph("Articles commonly implicated by power plans and one-lines may include: 90 (Introduction & Enforcement), "
                      "100 (Definitions), 110 (Requirements for Electrical Installations), 200 (Use and Identification of Grounded Conductors), "
                      "210 (Branch Circuits), 215 (Feeders), 220 (Branch-Circuit, Feeder, and Service Calculations), "
                      "225 (Outside Branch Circuits and Feeders), 230 (Services), 240 (Overcurrent Protection), "
                      "250 (Grounding and Bonding), 300 (Wiring Methods), 310 (Conductors), 408 (Switchboards, Switchgear, Panelboards), "
                      "450 (Transformers), 700/701/702 (Emergency/Legally Required/Optional Standby), 760 (Fire Alarm), 800+ (Communications) as applicable.")

    out_path = out_dir / _short_filename('summary', 'docx', session)
    doc.save(str(out_path))
    return out_path

@app.post("/export/build_zip")
def export_build_zip(payload: dict):
    """
    Finalizes a build:
      - Re-runs the selected task from 'payload' to ensure fresh outputs
      - Creates a summary DOCX capturing scope, ethics, PE checklist, and NEC references
      - Zips all generated artifacts for download
    Expected payload:
      {
        "intent": "one_line" | "power_plan" | "lighting_plan",
        "plan": {...},                   # structured JSON produced by plan_from_prompt (or customized)
        "outputs": ["dxf","pdf","csv"]   # desired formats
        "session": "optional-session-id"
      }
    """
    plan = payload.get("plan") or {}
    intent = (payload.get("intent") or "").lower()
    outputs = payload.get("outputs") or ["dxf","pdf"]
    session = payload.get("session")

    OUT.mkdir(parents=True, exist_ok=True)

    generated = []

    if intent == "one_line":
        req = OneLineRequest(**{
            "project": plan.get("project",""),
            "service_voltage": plan.get("service_voltage","480Y/277V"),
            "service_amperes": plan.get("service_amperes",2000),
            "panels": plan.get("panels",[]),
            "loads": plan.get("loads",[])
        })
        dxf_name = _short_filename('one_line', 'dxf', session)
        dxf_path = OUT / dxf_name
        generate_one_line_dxf(req, dxf_path)
        generated.append(dxf_name)
        if "pdf" in outputs:
            from app.export.pdf_from_dxf import dxf_to_pdf
            pdf_name = dxf_name.replace(".dxf",".pdf")
            dxf_to_pdf(dxf_path, OUT / pdf_name)
            generated.append(pdf_name)

    elif intent == "power_plan":
        req = PlanRequest(**{
            "project": plan.get("project",""),
            "rooms": plan.get("rooms",[]),
            "devices": plan.get("devices",[])
        })
        dxf_name = _short_filename('power_plan', 'dxf', session)
        dxf_path = OUT / dxf_name
        generate_power_plan_dxf(req, dxf_path)
        generated.append(dxf_name)
        if "pdf" in outputs:
            from app.export.pdf_from_dxf import dxf_to_pdf
            pdf_name = dxf_name.replace(".dxf",".pdf")
            dxf_to_pdf(dxf_path, OUT / pdf_name)
            generated.append(pdf_name)

    elif intent == "lighting_plan":
        req = PlanRequest(**{
            "project": plan.get("project",""),
            "rooms": plan.get("rooms",[]),
            "devices": plan.get("devices",[])
        })
        dxf_name = _short_filename('lighting_plan', 'dxf', session)
        dxf_path = OUT / dxf_name
        generate_lighting_plan_dxf(req, dxf_path)
        generated.append(dxf_name)
        if "pdf" in outputs:
            from app.export.pdf_from_dxf import dxf_to_pdf
            pdf_name = dxf_name.replace(".dxf",".pdf")
            dxf_to_pdf(dxf_path, OUT / pdf_name)
            generated.append(pdf_name)
    
    elif intent == "panel_schedule":
        # Use the /panel/ocr_to_excel endpoint logic
        pref = _session_prefix(session)
        BUCKET.mkdir(exist_ok=True)
        
        # Get image files from bucket
        image_files = [p.name for p in BUCKET.iterdir() 
                      if p.is_file() and p.suffix.lower() in [".jpg",".jpeg",".png",".tif",".tiff",".bmp"]
                      and (not pref or p.name.startswith(pref))]
        
        if image_files:
            all_lines = []
            for img_name in image_files:
                lines = ocr_image_to_lines(BUCKET / img_name)
                all_lines.extend(lines)
            
            circuits = parse_circuits_from_lines(all_lines, plan.get("number_of_ckts"))
            panel_specs_ocr = extract_panel_specs(all_lines)
        else:
            circuits = []
            panel_specs_ocr = {}
        
        # Merge OCR specs with user-provided specs (user specs take priority)
        panel_specs = panel_specs_ocr.copy()
        if plan.get("panel_specs"):
            panel_specs.update(plan["panel_specs"])
        
        # Use panel name from user input (chat/voice), then OCR, then default to MISSING
        panel_name = plan.get("panel_name") or panel_specs.get("panel_name", "MISSING")
        
        # Look for template
        template = find_template(BUCKET, pref)
        
        # Generate Excel (simplified - just copy template with parameter updates for now)
        xlsx_name = _short_filename('panel_schedule', 'xlsx', session)
        apply_template_to_data(circuits, panel_name, template, OUT / xlsx_name, panel_specs)
        generated.append(xlsx_name)

    # Optional CSV panel schedule if present in plan
    if "panel_schedule" in plan and isinstance(plan["panel_schedule"], dict):
        by_panel = plan["panel_schedule"]
        import csv, io
        csv_buf = io.StringIO()
        writer = csv.writer(csv_buf)
        writer.writerow(["Panel", "Load Name", "kVA"])
        for pnl, loads in by_panel.items():
            for ld in loads:
                writer.writerow([pnl, ld.get("name",""), ld.get("kva","")])
        csv_name = _short_filename('panel_schedule', 'csv', session)
        (OUT / csv_name).write_text(csv_buf.getvalue(), encoding="utf-8")
        generated.append(csv_name)

    # Create summary DOCX
    summary_path = _write_summary_docx(plan, OUT, session)
    generated.append(summary_path.name)

    # Zip it all
    zip_name = _short_filename('build', 'zip', session)
    zip_path = OUT / zip_name
    with _zipfile.ZipFile(zip_path, "w", _zipfile.ZIP_DEFLATED) as z:
        for name in generated:
            z.write(OUT / name, arcname=name)

    return {"message": "Build package ready.", "zip": zip_name, "artifacts": generated}


# ---- Session helpers ----
def _session_prefix(session: str|None) -> str:
    if not session:
        return ""
    # keep it filesystem-safe
    safe = "".join(ch for ch in session if ch.isalnum() or ch in ("-","_"))[:40]
    return f"{safe}__"

def _filter_session(files, session: str|None):
    pref = _session_prefix(session)
    if not pref:
        return files
    return [f for f in files if f.startswith(pref)]

def _short_filename(file_type: str, extension: str, session: str | None = None) -> str:
    """
    Generate short filename with optional session prefix.
    Format: {session}___{type}_{id}.{ext}
    Example: 3980c5ea__ps_a1b2c3.xlsx
    """
    # Type abbreviations
    type_map = {
        'one_line': 'ol',
        'power_plan': 'pp',
        'lighting_plan': 'lp',
        'panel_schedule': 'ps',
        'build': 'bld',
        'summary': 'sum'
    }
    
    prefix = type_map.get(file_type, file_type[:2])
    # Use 6 hex chars from UUID for uniqueness
    short_id = uuid.uuid4().hex[:6]
    
    # Build filename with session prefix if provided
    session_part = _session_prefix(session)
    filename = f"{session_part}{prefix}_{short_id}.{extension}"
    return filename


# ---- Panel OCR → Excel ----
from app.skills.ocr_panel import ocr_image_to_lines, parse_circuits_from_lines, extract_panel_specs
from app.utils.excel_template import find_template, apply_template_to_data
import openpyxl

@app.post("/panel/ocr_to_excel")
def panel_ocr_to_excel(payload: dict):
    """
    Create a panelboard schedule Excel from photos in the bucket.
    Automatically uses an Excel template if one is uploaded to the bucket.
    Template files should have 'template' in the filename (e.g., 'panelboard_template.xlsx')
    
    Payload:
      {
        "files": ["PB1_1.jpg","PB1_2.jpg"] | null to use all images in bucket for session
        "panel_name": "PB1",
        "number_of_ckts": 42 | null (18-80, must be even),
        "session": "optional-session-id",
        "use_template": true | false (default: true)
      }
    Requires Tesseract OCR installed locally.
    """
    session = payload.get("session")
    pref = _session_prefix(session)
    use_template = payload.get("use_template", True)
    BUCKET.mkdir(exist_ok=True)
    
    # Choose files
    files = payload.get("files")
    if not files:
        # use all images in bucket for this session
        cand = [p.name for p in BUCKET.iterdir() if p.is_file() and p.suffix.lower() in [".jpg",".jpeg",".png",".tif",".tiff",".bmp"]]
        files = [f for f in cand if not pref or f.startswith(pref)]
    if not files:
        raise HTTPException(400, "No image files found in bucket for this session. Upload photos first.")

    all_lines = []
    for name in files:
        path = BUCKET / name
        if not path.exists():
            raise HTTPException(400, f"Missing file: {name}")
        lines = ocr_image_to_lines(path)
        all_lines.extend(lines)

    panel_specs = extract_panel_specs(all_lines)
    
    # Get number_of_ckts from payload, panel specs, or default determination
    number_of_ckts = payload.get("number_of_ckts")
    if not number_of_ckts and 'number_of_ckts' in panel_specs:
        try:
            number_of_ckts = int(panel_specs['number_of_ckts'])
        except (ValueError, TypeError):
            number_of_ckts = None
    
    circuits = parse_circuits_from_lines(all_lines, number_of_ckts)
    
    # Use panel name from OCR extraction, default to MISSING if not found
    # Allow user override via payload if explicitly provided
    panel = payload.get("panel_name") or panel_specs.get("panel_name", "MISSING")
    
    # Store the determined number_of_ckts in panel_specs for template
    panel_specs['number_of_ckts'] = len(circuits)

    # Look for template
    template_path = None
    if use_template:
        template_path = find_template(BUCKET, pref)

    # Build Excel with template or basic format
    OUT.mkdir(parents=True, exist_ok=True)
    name = _short_filename('panel_schedule', 'xlsx', session)
    output_path = OUT / name
    
    # Use template-based generation
    apply_template_to_data(circuits, panel, template_path, output_path, panel_specs)
    
    # Also create raw OCR sheet for debugging
    wb = openpyxl.load_workbook(output_path)
    ws_raw = wb.create_sheet("Raw_OCR", 0)  # Insert as first sheet
    ws_raw.append(["Line #", "Text"])
    for i, ln in enumerate(all_lines, 1):
        ws_raw.append([i, ln])
    wb.save(output_path)
    
    return {
        "file": name, 
        "parsed_circuits": len(circuits), 
        "images_used": files,
        "template_used": template_path.name if template_path else None
    }
